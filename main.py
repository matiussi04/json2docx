from docx import Document
from bs4 import BeautifulSoup
from html2json.script import convert
import json
from docxtpl import DocxTemplate
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

styles = {
    "p": "paragrafo",
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 5",
}

align = {
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
}

OPTIONS = {
    "bold": False,
    "italic": False,
    "underline": False,
}


def is_html(text):
    soup = BeautifulSoup(text, "html.parser")
    return bool(soup.find())


def check_styles():
    doc_teste = Document("template.docx")
    styles_in_document = doc_teste.styles

    for style in styles_in_document:
        print(style.name)


def check_sections():
    doc_teste = Document("documento_final.docx")
    sections = doc_teste.sections

    for section in sections:
        print(section)


def data_table(data):
    table_data = data['values'][0]['values']

    matrix = []

    for row in table_data:
        row_values = []
        for cell in row['values']:
            row_values.append(cell['values'])
        matrix.append(row_values)

    return matrix


def render_values(paragraph, item, options=OPTIONS):
    for value in item:
        run = paragraph.add_run()
        if not paragraph.style.name.startswith("Heading"):
            run.italic = options["italic"]
            run.bold = options["bold"]
            run.underline = options["underline"]
        if isinstance(value, str):
            run.add_text(value)
        elif value["tag_name"] == "img":
            width = value["attributes"].get("width", "")
            height = value["attributes"].get("height", "")
            source = value["attributes"]["src"]
            if (width == ""):
                run.add_picture(source)
            else:
                width_in_pixels = float(width)
                height_in_pixels = float(height)
                run.add_picture(source, width=Inches(
                    width_in_pixels * 0.0138889), height=Inches(height_in_pixels * 0.0138889))
            paragraph.style = 'Normal'
        elif value["tag_name"] == "strong":
            run.bold = True
            new_options = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            render_values(paragraph, value["values"], new_options)
        elif value["tag_name"] == "em":
            run.italic = True
            new_options = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            render_values(paragraph, value["values"], new_options)
        elif value["tag_name"] == "span":
            run.underline = True
            new_options = {
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            render_values(paragraph, value["values"], new_options)


def process_list(subDocument, li):
    for item in li["values"]:
        p = subDocument.add_paragraph("Teste", style="List Paragraph")
        render_values(p, item["values"])


def process_table(subDocument, data):
    table = subDocument.add_table(
        rows=len(data), cols=len(data[0]), style="Table Grid")

    for row_index, row in enumerate(data):
        row_cells = table.rows[row_index].cells
        for index, cell_data in enumerate(row):
            p = row_cells[index].paragraphs[0]
            p.style = 'Normal'

            if isinstance(cell_data[0], str):
                render_values(p, cell_data)
            else:
                for index_item, item in enumerate(cell_data):
                    if index_item == 0:
                        render_values(p, item["values"])
                    else:
                        paragraph = row_cells[index].add_paragraph()
                        paragraph.style = 'Normal'
                        render_values(paragraph, item["values"])


def process_default(subDocument, item):
    paragraph = subDocument.add_paragraph()
    paragraph.style = item['classes'] if item['classes'] != "" else "paragrafo"
    render_values(paragraph, item['values'])


def process_items(doc, items):
    subDocument = doc.new_subdoc()

    for key in items["keys"]:
        if key["tag_name"] == "ul":
            process_list(subDocument, key)
        elif key["tag_name"] == "table":
            data = data_table(key)
            process_table(subDocument, data)
        else:
            process_default(subDocument, key)

    return subDocument


def json2docx(data):
    doc = DocxTemplate("template.docx")

    context = {}

    for field in data["fields"]:
        key = field["key"]
        value = field["value"]

        if type(value) is list:
            print('list')
        elif is_html(value):
            json_string = convert(value)
            items = json.loads(json_string)
            context.update({key: process_items(doc, items)})
        else:
            context[key] = value

    doc.render(context)

    doc.save("documento_final.docx")


if __name__ == '__main__':
    data = {
        "fields": [
            {
                "key": "titulo",
                "value": "TCC CONTROL: UM SISTEMA PARA GERENCIAMENTO DE TCC",
                "type": "string"
            },
            {
                "key": "coordenador",
                "value": "Lucas",
                "type": "string"
            },
            {
                "key": "aluno",
                "value": "Lucas",
                "type": "string"
            },
            {
                "key": "orientador",
                "value": "Felipe Pereira Perez",
                "type": "string"
            },
            {
                "key": "bibliografia",
                "value": [
                    "Bibliografia 1",
                    "Bibliografia 2",
                    "Bibliografia 3"
                ],
                "type": "ref"
            },
            {
                "key": "version",
                "value": [
                    {"created_at": "24/03/2023", "version": "1.0",
                        "description": "blablabla", "author": "Willian"},
                    {"created_at": "25/03/2023", "version": "1.1",
                        "description": "blablabla", "author": "Willian"},
                    {"created_at": "26/03/2023", "version": "1.2",
                        "description": "blablabla", "author": "Willian"},
                ],
                "type": "list_version"
            },
            {
                "key": "content",
                "value": """
<h1 class="Heading 1" style="font-size: 18px; text-transform: uppercase; font-weight: bold;">introdu&ccedil;&atilde;o</h1>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<h1 class="Heading 1" style="font-size: 18px; text-transform: uppercase; font-weight: bold;">conteudo</h1>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<h2 class="Heading 2" style="font-size: 18px; text-transform: uppercase; font-weight: normal;">CONTEudo 2</h2>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<h3 class="Heading 3" style="font-size: 18px; text-transform: capitalize; font-weight: bold;">Conteudo Terceiro</h3>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<h3 class="Heading 3" style="font-size: 18px; text-transform: capitalize; font-weight: bold;">Conteudo Terceiro Segundo</h3>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
<p>Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.</p>
    <p><img src="TCC Control.png" width="300" height="200"></img></p>
"""
            }
        ]
    }

    json2docx(data)
